"""Generates test CV fixtures. Run once after pulling, commit the result.

Usage:
    cd backend
    python tests/fixtures/_generate.py
"""

from pathlib import Path

from docx import Document
from reportlab.pdfgen import canvas

FIXTURES_DIR = Path(__file__).parent


def make_sample_pdf() -> None:
    path = FIXTURES_DIR / "cv_sample.pdf"
    c = canvas.Canvas(str(path))
    # Page 1
    c.drawString(100, 750, "Jean Dupont")
    c.drawString(100, 730, "Email: jean.dupont@example.com")
    c.drawString(100, 700, "EXPERIENCE PROFESSIONNELLE")
    c.drawString(100, 680, "Developpeur Python - Cinq ans")
    c.drawString(100, 660, "Architecte logiciel - Trois ans")
    c.showPage()
    # Page 2
    c.drawString(100, 750, "FORMATION")
    c.drawString(100, 730, "Master en informatique - Universite de Montreal")
    c.drawString(100, 710, "COMPETENCES")
    c.drawString(100, 690, "Python, FastAPI, PostgreSQL, Docker")
    c.save()


def make_sample_docx() -> None:
    path = FIXTURES_DIR / "cv_sample.docx"
    doc = Document()
    doc.add_paragraph("Marie Tremblay")
    doc.add_paragraph("Email: marie.tremblay@example.com")
    doc.add_paragraph("")
    doc.add_paragraph("EXPERIENCE PROFESSIONNELLE")
    doc.add_paragraph("Data scientist - Quatre ans")
    doc.add_paragraph("")
    doc.add_paragraph("FORMATION")
    doc.add_paragraph("Doctorat en statistiques")
    doc.add_paragraph("")
    doc.add_paragraph("COMPETENCES")
    doc.add_paragraph("Python, R, scikit-learn, TensorFlow")
    doc.save(str(path))


def make_corrupt_pdf() -> None:
    path = FIXTURES_DIR / "cv_corrupt.pdf"
    path.write_bytes(b"this is not a real PDF, just random bytes pretending to be one.")


def make_image_only_pdf() -> None:
    """Empty PDF (single blank page, no text). pdfplumber returns empty string."""
    path = FIXTURES_DIR / "cv_image_only.pdf"
    c = canvas.Canvas(str(path))
    c.showPage()  # blank page
    c.save()


if __name__ == "__main__":
    make_sample_pdf()
    make_sample_docx()
    make_corrupt_pdf()
    make_image_only_pdf()
    print("Generated fixtures in", FIXTURES_DIR)
